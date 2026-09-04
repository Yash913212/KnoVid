"""Document parser for uploaded files.

Extracts text, headings, and outlines from PDF (.pdf), Word (.docx),
Markdown (.md), Plain text (.txt), CSV (.csv), and JSON (.json).
Converts document structure into timed segments and semantic chapters.
"""
import json
import logging
from pathlib import Path

from pydantic import BaseModel

from app.schemas.schemas import ChapterOut, SegmentOut
from app.services.web_scraper import WORDS_PER_SECOND, _clean_text

logger = logging.getLogger(__name__)

DOCUMENT_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md", ".json", ".csv"}
AUDIO_EXTENSIONS = {".mp3", ".wav", ".m4a", ".aac", ".flac", ".ogg", ".opus", ".wma"}
VIDEO_EXTENSIONS = {".mp4", ".mkv", ".webm", ".mov", ".avi", ".flv", ".wmv", ".m4v"}


def is_document_file(file_path: Path | str) -> bool:
    """Check if the given path is a document (PDF, DOCX, TXT, MD, etc.)."""
    return Path(file_path).suffix.lower() in DOCUMENT_EXTENSIONS


def is_audio_file(file_path: Path | str) -> bool:
    """Check if the given path is an audio-only file."""
    return Path(file_path).suffix.lower() in AUDIO_EXTENSIONS


class ParsedDocument(BaseModel):
    title: str
    filePath: str
    duration: float
    language: str
    segments: list[SegmentOut]
    chapters: list[ChapterOut]
    full_text: str


def parse_uploaded_document(file_path: Path) -> ParsedDocument:
    """Parse any uploaded document file into structured segments and chapters."""
    ext = file_path.suffix.lower()
    title = file_path.stem.replace("-", " ").replace("_", " ").title()

    if ext == ".pdf":
        return _parse_pdf_file(file_path, title)
    elif ext in (".docx", ".doc"):
        return _parse_docx_file(file_path, title)
    elif ext in (".json", ".csv"):
        return _parse_data_file(file_path, title, ext)
    else:
        # Default text / markdown
        return _parse_text_file(file_path, title)


def _parse_pdf_file(file_path: Path, default_title: str) -> ParsedDocument:
    import pypdf

    with open(file_path, "rb") as f:
        reader = pypdf.PdfReader(f)
        num_pages = len(reader.pages)

        title = default_title
        if reader.metadata and reader.metadata.title:
            extracted_title = str(reader.metadata.title).strip()
            if extracted_title:
                title = extracted_title

        segments: list[SegmentOut] = []
        chapters: list[ChapterOut] = []
        current_time = 0.0

        for page_idx, page in enumerate(reader.pages):
            page_text = _clean_text(page.extract_text() or "")
            if not page_text:
                continue

            page_start = current_time
            paragraphs = [p.strip() for p in page_text.split("\n\n") if len(p.strip()) > 10]
            if not paragraphs:
                paragraphs = [page_text]

            for p in paragraphs:
                words = p.split()
                chunk_size = 70
                for i in range(0, len(words), chunk_size):
                    chunk_words = words[i:i + chunk_size]
                    chunk_text = " ".join(chunk_words)
                    duration = max(3.0, len(chunk_words) / WORDS_PER_SECOND)

                    segments.append(SegmentOut(
                        start=round(current_time, 2),
                        end=round(current_time + duration, 2),
                        speaker=f"Page {page_idx + 1}",
                        language="en",
                        text=chunk_text,
                        confidence=1.0,
                    ))
                    current_time += duration

            chapters.append(ChapterOut(
                id=f"ch_{page_idx}",
                title=f"Page {page_idx + 1}",
                start=round(page_start, 2),
                end=round(current_time, 2),
                summary=f"Key points from page {page_idx + 1} of {num_pages}",
                keywords=[],
            ))

        total_duration = max(10.0, round(current_time, 2))
        full_text = " ".join(s.text for s in segments)

        return ParsedDocument(
            title=title[:200],
            filePath=str(file_path),
            duration=total_duration,
            language="en",
            segments=segments,
            chapters=chapters,
            full_text=full_text,
        )


def _parse_docx_file(file_path: Path, default_title: str) -> ParsedDocument:
    import docx

    doc = docx.Document(str(file_path))
    segments: list[SegmentOut] = []
    chapters: list[ChapterOut] = []

    current_section = default_title
    chapter_start_time = 0.0
    current_time = 0.0
    chapter_idx = 0
    chapter_keywords: list[str] = []

    for para in doc.paragraphs:
        text = _clean_text(para.text)
        if not text or len(text) < 10:
            continue

        style_name = (para.style.name or "").lower() if para.style else ""
        is_heading = "heading" in style_name or "title" in style_name

        if is_heading:
            if current_time > chapter_start_time:
                chapters.append(ChapterOut(
                    id=f"ch_{chapter_idx}",
                    title=current_section[:80],
                    start=round(chapter_start_time, 2),
                    end=round(current_time, 2),
                    summary=f"Discussion of {current_section}",
                    keywords=chapter_keywords[:5],
                ))
                chapter_idx += 1
                chapter_keywords = []

            current_section = text
            chapter_start_time = current_time
            duration = max(3.0, len(text.split()) / WORDS_PER_SECOND)
            segments.append(SegmentOut(
                start=round(current_time, 2),
                end=round(current_time + duration, 2),
                speaker="Heading",
                language="en",
                text=f"## {text}",
                confidence=1.0,
            ))
            current_time += duration
            continue

        words = text.split()
        chunk_size = 65
        for i in range(0, len(words), chunk_size):
            chunk_words = words[i:i + chunk_size]
            chunk_text = " ".join(chunk_words)
            duration = max(3.0, len(chunk_words) / WORDS_PER_SECOND)

            segments.append(SegmentOut(
                start=round(current_time, 2),
                end=round(current_time + duration, 2),
                speaker=current_section[:24],
                language="en",
                text=chunk_text,
                confidence=1.0,
            ))
            current_time += duration

            for w in chunk_words:
                if len(w) > 6 and w.isalpha():
                    chapter_keywords.append(w.lower())

    if current_time > chapter_start_time:
        chapters.append(ChapterOut(
            id=f"ch_{chapter_idx}",
            title=current_section[:80],
            start=round(chapter_start_time, 2),
            end=round(current_time, 2),
            summary=f"Section {current_section}",
            keywords=chapter_keywords[:5],
        ))

    total_duration = max(10.0, round(current_time, 2))
    full_text = " ".join(s.text for s in segments)

    return ParsedDocument(
        title=default_title[:200],
        filePath=str(file_path),
        duration=total_duration,
        language="en",
        segments=segments,
        chapters=chapters,
        full_text=full_text,
    )


def _parse_text_file(file_path: Path, default_title: str) -> ParsedDocument:
    raw_content = ""
    for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            raw_content = file_path.read_text(encoding=enc)
            break
        except Exception:
            continue

    if not raw_content:
        raw_content = f"Document: {default_title}"

    lines = raw_content.splitlines()
    segments: list[SegmentOut] = []
    chapters: list[ChapterOut] = []

    current_section = default_title
    chapter_start_time = 0.0
    current_time = 0.0
    chapter_idx = 0

    accumulated: list[str] = []

    def flush_accumulated():
        nonlocal current_time
        if not accumulated:
            return
        block = " ".join(accumulated)
        accumulated.clear()
        words = block.split()
        chunk_size = 65
        for i in range(0, len(words), chunk_size):
            chunk_words = words[i:i + chunk_size]
            chunk_text = " ".join(chunk_words)
            duration = max(3.0, len(chunk_words) / WORDS_PER_SECOND)
            segments.append(SegmentOut(
                start=round(current_time, 2),
                end=round(current_time + duration, 2),
                speaker=current_section[:24],
                language="en",
                text=chunk_text,
                confidence=1.0,
            ))
            current_time += duration

    for line in lines:
        s_line = line.strip()
        if not s_line:
            flush_accumulated()
            continue

        # Markdown headings (# Heading)
        if s_line.startswith("#"):
            flush_accumulated()
            heading_text = s_line.lstrip("#").strip()
            if current_time > chapter_start_time:
                chapters.append(ChapterOut(
                    id=f"ch_{chapter_idx}",
                    title=current_section[:80],
                    start=round(chapter_start_time, 2),
                    end=round(current_time, 2),
                    summary=f"Section discussing {current_section}",
                    keywords=[],
                ))
                chapter_idx += 1

            current_section = heading_text or default_title
            chapter_start_time = current_time

            duration = max(3.0, len(heading_text.split()) / WORDS_PER_SECOND)
            segments.append(SegmentOut(
                start=round(current_time, 2),
                end=round(current_time + duration, 2),
                speaker="Section",
                language="en",
                text=s_line,
                confidence=1.0,
            ))
            current_time += duration
            continue

        accumulated.append(s_line)

    flush_accumulated()

    if current_time > chapter_start_time:
        chapters.append(ChapterOut(
            id=f"ch_{chapter_idx}",
            title=current_section[:80],
            start=round(chapter_start_time, 2),
            end=round(current_time, 2),
            summary=f"Conclusion of {current_section}",
            keywords=[],
        ))

    total_duration = max(10.0, round(current_time, 2))
    full_text = " ".join(s.text for s in segments)

    return ParsedDocument(
        title=default_title[:200],
        filePath=str(file_path),
        duration=total_duration,
        language="en",
        segments=segments,
        chapters=chapters,
        full_text=full_text,
    )


def _parse_data_file(file_path: Path, default_title: str, ext: str) -> ParsedDocument:
    raw_content = file_path.read_text(encoding="utf-8", errors="replace")
    segments: list[SegmentOut] = []
    chapters: list[ChapterOut] = []
    current_time = 0.0

    if ext == ".json":
        try:
            data = json.loads(raw_content)
            formatted = json.dumps(data, indent=2)
        except Exception:
            formatted = raw_content
    else:
        formatted = raw_content

    lines = formatted.splitlines()
    chunk_size = 20
    for idx in range(0, len(lines), chunk_size):
        chunk = "\n".join(lines[idx:idx + chunk_size])
        duration = 15.0
        segments.append(SegmentOut(
            start=round(current_time, 2),
            end=round(current_time + duration, 2),
            speaker=f"Records {idx + 1}-{min(idx + chunk_size, len(lines))}",
            language="en",
            text=chunk,
            confidence=1.0,
        ))
        current_time += duration

    chapters.append(ChapterOut(
        id="ch_0",
        title=f"{ext.upper().lstrip('.')} Data Overview",
        start=0.0,
        end=round(current_time, 2),
        summary=f"Structured {ext} dataset with {len(lines)} lines",
        keywords=[],
    ))

    return ParsedDocument(
        title=default_title[:200],
        filePath=str(file_path),
        duration=max(10.0, round(current_time, 2)),
        language="en",
        segments=segments,
        chapters=chapters,
        full_text=raw_content[:20000],
    )
